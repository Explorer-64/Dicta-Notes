import logging
import time
from fastapi import APIRouter, HTTPException, Depends, Response

logger = logging.getLogger("dicta.export")
from app.libs.storage_manager import get_json as storage_get_json
from typing import Dict, Any
import traceback
import io # Added for in-memory file handling
from docx import Document # Added for DOCX generation

# ReportLab Imports for PDF Generation
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.lib.units import inch
from reportlab.lib.colors import grey
from reportlab.pdfgen import canvas

# Import models (ensure these exist in your models file)
from app.apis.models import ExportSessionRequest, Session, TranscriptionResponse, TranscriptionSegment, Speaker

# Import dependencies & helpers
from app.apis.helpers import get_current_user, get_session_key # Need session key helper
from app.apis.firebase import get_firestore_db # Need Firestore access
from docx.shared import Inches, Pt # Added for styling
from docx.enum.text import WD_ALIGN_PARAGRAPH # Added for alignment

# Define user dependency
user_dependency = Depends(get_current_user)

router = APIRouter(prefix="/export", tags=["Export"])

def _format_timestamp(seconds: float) -> str:
    """Formats seconds into HH:MM:SS string."""
    if seconds is None or seconds < 0:
        # Use a more distinct placeholder if time is unknown
        return "[--:--:--]"
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    return f"[{hours:02d}:{minutes:02d}:{secs:02d}]"


# Helper to add styled paragraphs
def add_styled_paragraph(document, text, style_name=None, bold=False, italic=False, size=None, alignment=None):
    p = document.add_paragraph(text)
    if style_name:
        p.style = style_name
    run = p.runs[0]
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if alignment:
        p.alignment = alignment
    return p


@router.post("/export/{session_id}", response_class=Response) # <-- Restored path parameter
def export_session(req: ExportSessionRequest, current_user: Dict[str, Any] = user_dependency):
    session_id = req.session_id # <-- Still get session_id from body
    logger.debug("Export request for session_id: %s", session_id)
    """Exports a session transcript based on selected format and options."""
    logger.info("Exporting session %s in format %s", session_id, req.format)
    user_id = current_user.get("uid")

    try:
        # TIER GATE: Check if user can export in this format
        from app.libs.tier_management import TierManager
        tier_manager = TierManager()
        
        can_export, gate_message = tier_manager.can_export_format(user_id, req.format)
        
        if not can_export:
            logger.warning("Export format '%s' blocked for %s: %s", req.format, user_id, gate_message)
            raise HTTPException(
                status_code=403,
                detail=gate_message or f"Export format '{req.format}' is not available on your plan. Upgrade to Individual or higher."
            )
        logger.debug("Export format '%s' allowed for %s", req.format, user_id)
        
        # --- Fetch session data (Firestore first, then fallback) ---
        firestore_db = get_firestore_db()
        session_data = None
        # using_firestore = False # REMOVED

        if firestore_db:
            try:
                doc_ref = firestore_db.collection('sessions').document(session_id)
                session_doc = doc_ref.get()
                if session_doc.exists:
                    session_data = session_doc.to_dict()
                    # Security check
                    if session_data.get("userId") != user_id and not session_data.get("companyId"): # Basic check
                         raise HTTPException(status_code=403, detail="Not authorized to access this session")
                else:
                    logger.debug("Session %s not found in Firestore for export", session_id)
            except Exception as firestore_error:
                logger.warning("Error accessing Firestore for export: %s", firestore_error)
                # Continue to fallback

        if not session_data:
             try:
                logger.info("Trying fallback storage for export of session %s", session_id)
                session_key = get_session_key(session_id)
                session_data = storage_get_json(session_key)
                if not session_data:
                    raise HTTPException(status_code=404, detail=f"Session {session_id} not found")
                # Security check
                if session_data.get("userId") != user_id:
                     raise HTTPException(status_code=403, detail="Not authorized to access this session")
             except Exception as storage_error:
                logger.warning("Error retrieving session %s from storage for export: %s", session_id, storage_error)
                if not isinstance(storage_error, HTTPException):
                     raise HTTPException(status_code=404, detail=f"Session {session_id} not found") from storage_error
                else:
                     raise storage_error # Re-raise specific HTTP exceptions

        # --- Prepare content (basic text for now) ---
        transcript_text = session_data.get("full_text", "")
        title = session_data.get("title", "transcript")
        client_name = session_data.get("client_name") # Get freelancer metadata
        project_name = session_data.get("project_name")

        # TODO: Apply formatting options from req.options (header, footer, timestamps, template, etc.)

        # --- Generate content based on format ---
        file_content = b"" # Initialize as bytes
        content_type = "text/plain"
        file_extension = "txt"

        if req.format == 'docx':
            try:
                document = Document()
                opts = req.options # Shortcut for options
                template_name = opts.professionalTemplate or 'standard'

                # --- Setup Styles (Basic example) ---
                # You might want to define styles based on templates here
                # e.g., document.styles['Normal'].font.name = 'Calibri'

                # --- Add Header --- 
                if opts.includeHeader and opts.headerText:
                    section = document.sections[0]
                    header = section.header
                    if not header.paragraphs:
                        header.add_paragraph()
                    header_paragraph = header.paragraphs[0]
                    header_paragraph.text = opts.headerText
                    header_paragraph.style = document.styles['Header']
                    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT # Example alignment
                    # Add branding space in header if requested
                    if opts.includeBranding:
                         header.add_paragraph("[Branding Space]").alignment = WD_ALIGN_PARAGRAPH.RIGHT

                # --- Add Footer --- 
                if opts.includeFooter and opts.footerText:
                    section = document.sections[0]
                    footer = section.footer
                    if not footer.paragraphs:
                        footer.add_paragraph()
                    footer_paragraph = footer.paragraphs[0]
                    footer_paragraph.text = opts.footerText
                    footer_paragraph.style = document.styles['Footer']
                    # Example: Add page number (requires more advanced setup)
                    # add_page_number(footer.paragraphs[0]) # Placeholder for potential future function

                # --- Add Document Title --- 
                # Template specific titles
                title_text = f"Transcript: {session_data.get('title', 'Untitled Meeting')}"
                if template_name == 'modern':
                    title_text = f"SESSION RECORD: {session_data.get('title', 'Untitled').upper()}"
                elif template_name == 'legal':
                     title_text = f"Official Record - Session ID: {session_id}"
                elif template_name == 'minimal':
                    title_text = session_data.get('title', 'Session Transcript')
                    
                add_styled_paragraph(document, title_text, style_name='Heading 1', size=16, bold=True)
                document.add_paragraph() # Spacing

                # --- Add Client/Project Info --- 
                if opts.includeClientProjectInfo: # Check the actual field name used
                    info_added = False
                    if client_name:
                        add_styled_paragraph(document, f"Client: {client_name}", bold=True)
                        info_added = True
                    if project_name:
                        add_styled_paragraph(document, f"Project: {project_name}", bold=True)
                        info_added = True
                    if info_added:
                         document.add_paragraph() # Add spacing if info was added

                # --- Add Content with Timestamps --- 
                segments = session_data.get('segments', [])
                timestamp_freq = opts.timestampFrequency
                last_speaker = None
                last_minute_marker = -1
                last_5min_marker = -1

                if not segments and transcript_text: # Check transcript_text if segments missing
                    # Fallback if segments are missing, use full_text
                    add_styled_paragraph(document, transcript_text)
                else:
                    for i, segment in enumerate(segments):
                        speaker = segment.get('speaker', 'Unknown Speaker')
                        text = segment.get('text', '').strip()
                        start_time = segment.get('start_time') # Float or None
                        
                        if not text: # Skip empty segments
                            continue

                        # --- Minute Markers --- 
                        current_minute = int(start_time // 60) if start_time is not None else -1
                        if timestamp_freq == 'every-minute' and current_minute > last_minute_marker:
                             add_styled_paragraph(document, f"--- {_format_timestamp(current_minute * 60)} --- ", bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                             last_minute_marker = current_minute
                        elif timestamp_freq == 'every-5-minutes' and current_minute // 5 > last_5min_marker:
                            marker_time = (current_minute // 5) * 5 * 60
                            add_styled_paragraph(document, f"--- {_format_timestamp(marker_time)} --- ", bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                            last_5min_marker = current_minute // 5
                            
                        # --- Speaker Line --- 
                        prefix = ""
                        if timestamp_freq == 'every-speaker' and start_time is not None:
                            prefix = f"{_format_timestamp(start_time)} "

                        # Add speaker name only if it changes or it's the first segment
                        speaker_line = f"{prefix}{speaker}:"
                        if speaker != last_speaker or i == 0:
                             add_styled_paragraph(document, speaker_line, bold=True)
                             last_speaker = speaker
                        else:
                             # If same speaker, maybe just add timestamp if freq is 'every-speaker'?
                             if timestamp_freq == 'every-speaker':
                                 add_styled_paragraph(document, f"{prefix}", bold=True) # Just timestamp
                             else: 
                                 # Or maybe indent subsequent lines? For simplicity, just add text.
                                 pass 
                                 
                        # Add the actual text
                        add_styled_paragraph(document, text)
                        
                        # Add small space between different speakers? (Optional)
                        # if i + 1 < len(segments) and segments[i+1].get('speaker') != speaker:
                        #    document.add_paragraph() 
                            

                # --- Add Notes --- 
                if opts.includeNotes:
                    notes = session_data.get('notes') # Get notes from session data
                    if notes: # Only add section if notes exist
                        document.add_paragraph() # Add space before notes
                        add_styled_paragraph(document, "Notes", style_name='Heading 2', bold=True)
                        add_styled_paragraph(document, notes) # Add the notes content

                # Save document to memory
                mem_file = io.BytesIO()
                document.save(mem_file)
                mem_file.seek(0)
                file_content = mem_file.read()

                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                file_extension = "docx"
            except ImportError as import_error:
                logger.error("ImportError generating DOCX for %s: %s. Is python-docx installed?", session_id, import_error)
                raise HTTPException(status_code=500, detail="Server configuration error: Required library missing.") from import_error
            except Exception as docx_error:
                logger.error("Error generating DOCX for %s: %s", session_id, docx_error)
                logger.debug("%s", traceback.format_exc())
                raise HTTPException(status_code=500, detail=f"Failed to generate DOCX file: {docx_error}") from docx_error

        elif req.format == 'pdf':
            try:
                mem_file = io.BytesIO()
                # Create a PDF document using SimpleDocTemplate for flowing elements
                doc = SimpleDocTemplate(mem_file, pagesize=(8.5*inch, 11*inch),
                                        leftMargin=0.75*inch, rightMargin=0.75*inch,
                                        topMargin=0.75*inch, bottomMargin=0.75*inch)
                styles = getSampleStyleSheet()
                story = [] # List to hold flowable elements for ReportLab
                opts = req.options # Shortcut for options
                template_name = opts.professionalTemplate or 'standard'

                # --- Add Header/Footer (using page templates is more robust, 
                #     but complex. For simplicity, we might add static text at top/bottom
                #     or handle via canvas events if needed later)
                # Basic Header placeholder
                if opts.includeHeader and opts.headerText:
                    story.append(Paragraph(opts.headerText, styles['Normal']))
                    if opts.includeBranding:
                         story.append(Paragraph("[Branding Space]", styles['Normal'])) # Simple text for now
                    story.append(Spacer(1, 0.2*inch))

                # --- Add Document Title ---
                title_text = f"Transcript: {session_data.get('title', 'Untitled Meeting')}"
                # Adjust title based on template (similar to DOCX)
                # ... (add template specific title logic here if needed)
                story.append(Paragraph(title_text, styles['h1']))
                story.append(Spacer(1, 0.1*inch))

                # --- Add Client/Project Info ---
                if opts.includeClientProjectInfo:
                    info_added = False
                    if client_name:
                        story.append(Paragraph(f"<b>Client:</b> {client_name}", styles['Normal']))
                        info_added = True
                    if project_name:
                        story.append(Paragraph(f"<b>Project:</b> {project_name}", styles['Normal']))
                        info_added = True
                    if info_added:
                        story.append(Spacer(1, 0.1*inch))

                # --- Add Content with Timestamps ---
                segments = session_data.get('segments', [])
                timestamp_freq = opts.timestampFrequency
                last_speaker = None
                last_minute_marker = -1
                last_5min_marker = -1
                body_style = styles['BodyText']
                # body_style.spaceAfter = 6 # Remove default space after all body text
                speaker_style = styles['BodyText'] # Style for speaker line
                speaker_style.spaceBefore = 10 # Add space BEFORE speaker line
                text_style = styles['BodyText'] # Style for the actual text
                text_style.spaceAfter = 6 # Add space AFTER the text block

                if not segments and transcript_text:
                    # If no segments but we have full text, split by newlines and add each as a separate paragraph
                    # This preserves the original formatting of the transcript
                    for line in transcript_text.split('\n'):
                        if line.strip():
                            story.append(Paragraph(line.replace('<', '&lt;').replace('>', '&gt;'), body_style))
                        else:
                            # Add a small spacer for empty lines to preserve paragraph spacing
                            story.append(Spacer(1, 0.1*inch))
                else:
                    for i, segment in enumerate(segments):
                        speaker = segment.get('speaker', 'Unknown Speaker')
                        text = segment.get('text', '').strip()
                        start_time = segment.get('start_time')

                        if not text:
                            continue

                        # --- Minute Markers --- 
                        current_minute = int(start_time // 60) if start_time is not None else -1
                        if timestamp_freq == 'every-minute' and current_minute > last_minute_marker:
                            marker_text = f"--- {_format_timestamp(current_minute * 60)} ---"
                            marker_style = styles['Normal']
                            marker_style.alignment = TA_CENTER
                            marker_style.textColor = grey
                            story.append(Spacer(1, 0.1*inch))
                            story.append(Paragraph(f"<i><b>{marker_text}</b></i>", marker_style))
                            story.append(Spacer(1, 0.1*inch))
                            last_minute_marker = current_minute
                        elif timestamp_freq == 'every-5-minutes' and current_minute // 5 > last_5min_marker:
                            marker_time = (current_minute // 5) * 5 * 60
                            marker_text = f"--- {_format_timestamp(marker_time)} ---"
                            marker_style = styles['Normal']
                            marker_style.alignment = TA_CENTER
                            marker_style.textColor = grey
                            story.append(Spacer(1, 0.1*inch))
                            story.append(Paragraph(f"<i><b>{marker_text}</b></i>", marker_style))
                            story.append(Spacer(1, 0.1*inch))
                            last_5min_marker = current_minute // 5

                        # --- Speaker Line --- 
                        prefix = ""
                        if timestamp_freq == 'every-speaker' and start_time is not None:
                            prefix = f"{_format_timestamp(start_time)} "

                        speaker_line = f"<b>{prefix}{speaker}:</b>"
                        if speaker != last_speaker or i == 0:
                            # Apply specific speaker style
                            story.append(Paragraph(speaker_line, speaker_style))
                            last_speaker = speaker
                        else:
                             if timestamp_freq == 'every-speaker':
                                 # Apply specific speaker style (even for just timestamp)
                                 story.append(Paragraph(f"<b>{prefix}</b>", speaker_style))
                             else:
                                 pass # No speaker line needed if same speaker and no per-speaker timestamp

                        # Add the actual text (escape HTML potentially harmful characters)
                        # Apply specific text style and ensure line breaks are preserved
                        # The text_style already has spaceAfter set to 6, which adds vertical spacing between paragraphs
                        story.append(Paragraph(text.replace('<', '&lt;').replace('>', '&gt;'), text_style))
                
                # --- Add Notes ---
                if opts.includeNotes:
                    notes_html = session_data.get('notes') # Get notes (could be HTML)
                    if notes_html:
                        story.append(Spacer(1, 0.2*inch))
                        story.append(Paragraph("<b>Notes</b>", styles['h2']))
                        # ReportLab's Paragraph can handle basic HTML-like tags
                        # For complex HTML, need external libraries like xhtml2pdf or rlextra
                        # Simple approach: Add as a paragraph, basic tags might work.
                        # Escape actual < > to avoid issues
                        safe_notes = notes_html.replace('<', '&lt;').replace('>', '&gt;')
                        # Replace common tags with ReportLab equivalents if needed
                        # safe_notes = safe_notes.replace('<p>', '').replace('</p>', '<br/>') # Example
                        story.append(Paragraph(safe_notes, styles['Normal']))
                        
                # --- Build the PDF --- 
                doc.build(story)
                mem_file.seek(0)
                file_content = mem_file.read()

                content_type = "application/pdf"
                file_extension = "pdf"
            
            except ImportError as import_error:
                logger.error("ImportError generating PDF for %s: %s. Is reportlab installed?", session_id, import_error)
                raise HTTPException(status_code=500, detail="Server configuration error: Required library missing.") from import_error
            except Exception as pdf_error:
                logger.error("Error generating PDF for %s: %s", session_id, pdf_error)
                logger.debug("%s", traceback.format_exc())
                raise HTTPException(status_code=500, detail=f"Failed to generate PDF file: {pdf_error}") from pdf_error

        # TODO: Add elif branches for other formats (json, md)
        else: # Default to plain text
            file_content = transcript_text.encode('utf-8') # Encode string to bytes
            content_type = "text/plain"
            file_extension = "txt"

        # Construct filename
        filename_parts = [title]
        if client_name:
            filename_parts.append(client_name)
        if project_name:
            filename_parts.append(project_name)
        filename_parts.append(session_id)
        filename = f"{'_'.join(filename_parts)}.{file_extension}"
        # Sanitize filename (simple example)
        filename = filename.replace(" ", "_").replace("/", "-")

        # Set headers for file download
        headers = {
            'Content-Disposition': f'attachment; filename="{filename}"'
        }

        return Response(content=file_content, media_type=content_type, headers=headers)

    except HTTPException as e:
        raise e # Re-raise HTTP exceptions
    except Exception as e:
        logger.error("Error exporting session %s: %s", session_id, e)
        logger.debug("%s", traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Internal server error during export: {str(e)}") from e # Use 'from e'
